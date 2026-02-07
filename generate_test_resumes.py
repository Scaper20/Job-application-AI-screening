
import os
from reportlab.pdfgen import canvas
import docx

def create_pdf(filename, name, email, phone, education, skills, age, experience):
    c = canvas.Canvas(filename)
    y = 800
    c.drawString(100, y, f"Name: {name}"); y -= 20
    c.drawString(100, y, f"Email: {email}"); y -= 20
    c.drawString(100, y, f"Phone: {phone}"); y -= 20
    c.drawString(100, y, f"Age: {age}"); y -= 20
    c.drawString(100, y, f"Experience: {experience} years"); y -= 20
    c.drawString(100, y, f"Education: {education}"); y -= 20
    c.drawString(100, y, f"Skills: {skills}"); y -= 20
    c.save()

def create_docx(filename, name, email, phone, education, skills, age, experience):
    doc = docx.Document()
    doc.add_paragraph(f"Name: {name}")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")
    doc.add_paragraph(f"Age: {age}")
    doc.add_paragraph(f"Experience: {experience} years")
    doc.add_paragraph(f"Education: {education}")
    doc.add_paragraph(f"Skills: {skills}")
    doc.save(filename)

if __name__ == "__main__":
    os.makedirs("test_data", exist_ok=True)
    
    # Resume 1 (PDF)
    create_pdf("test_data/resume1.pdf", "John Doe", "johndoe@example.com", "123-456-7890", 
               "Bachelors in Computer Science", "Python, Machine Learning, SQL", 25, 3)
    
    # Resume 2 (DOCX)
    create_docx("test_data/resume2.docx", "Jane Smith", "janesmith@test.com", "987-654-3210", 
                "Masters in Data Science", "Java, Tableau, Communication", 29, 5)
    
    # Resume 3 (PDF)
    create_pdf("test_data/resume3.pdf", "Alice Brown", "alice@example.com", "555-123-4567", 
               "PhD in AI", "Deep Learning, PyTorch, Research", 32, 8)
    
    # Resume 4 (DOCX)
    create_docx("test_data/resume4.docx", "Bob White", "bob@test.com", "444-987-6543", 
                "Diploma in IT", "Networking, Linux, Support", 22, 1)
                
    # Resume 5 (PDF)
    create_pdf("test_data/resume5.pdf", "Charlie Green", "charlie@example.com", "333-555-7777", 
               "Bachelors in Engineering", "Project Management, Agile, Jira", 27, 4)

    print("5 Test resumes created (3 PDF, 2 DOCX).")

